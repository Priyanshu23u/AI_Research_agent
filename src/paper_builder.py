import os
import re
import pandas as pd
from typing import Dict
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch


def clean(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def _authors_to_str(authors_val) -> str:
    if isinstance(authors_val, list):
        return ", ".join([clean(a) for a in authors_val if a and str(a).strip()])
    return clean(str(authors_val))


def format_papers_bullets(rows: pd.DataFrame, max_chars_per_abstract: int = 1200) -> str:
    lines = []
    for _, r in rows.iterrows():
        abs_text = clean(r.get("abstract", ""))[:max_chars_per_abstract]
        link = r.get("link_abs") or r.get("link_pdf", "")
        authors = _authors_to_str(r.get("authors", ""))
        year = clean(str(r.get("year", "")))
        line = f"- {r.get('title','')} — {authors} — {year} — {link}\n  Abstract: {abs_text}"
        lines.append(line)
    return "\n".join(lines)


def make_prompt(topic: str, papers_df: pd.DataFrame, section_name: str, requirements: str) -> str:
    papers_bullets = format_papers_bullets(papers_df) if papers_df is not None and not papers_df.empty else "None"
    prompt = (
        "You are an AI research writer. Based ONLY on the provided papers (titles, abstracts, authors, links), "
        "write the requested section in a formal academic tone. Do NOT fabricate references. Use only given titles/links.\n\n"
        f"Topic: {topic}\n\n"
        "PAPERS (Title — Authors — Year — Link — Abstract):\n"
        f"{papers_bullets}\n\n"
        "Write the following section now:\n"
        f"- Section: {section_name}\n"
        f"- Requirements: {requirements}\n\n"
        f"Begin the {section_name}:\n"
    )
    return prompt


def refs_markdown(rows: pd.DataFrame) -> str:
    if rows is None or rows.empty:
        return ""
    refs = []
    for i, (_, r) in enumerate(rows.iterrows(), start=1):
        title = clean(r.get("title", ""))
        authors = _authors_to_str(r.get("authors", ""))
        year = r.get("year", "") or "n.d."
        link = r.get("link_abs") or r.get("link_pdf", "")
        refs.append(f"[{i}] {title} — {authors} ({year}) {link}")
    return "\n".join(refs)


def assemble_markdown(sections: Dict[str, str], refs_md: str, title: str) -> str:
    md = f"# {title}\n\n"
    order = [
        "Abstract", "Introduction", "Literature Review",
        "Key Findings", "Results & Discussion", "Conclusion", "Future Works"
    ]
    for sec in order:
        content = sections.get(sec, "").strip()
        if content:
            md += f"## {sec}\n\n{content}\n\n"
    if refs_md.strip():
        md += "## References\n\n" + refs_md + "\n"
    return md


def write_docx(path: str, title: str, sections: Dict[str, str], references_md: str):
    doc = Document()
    doc.add_heading(title, 0)

    order = [
        "Abstract", "Introduction", "Literature Review",
        "Key Findings", "Results & Discussion", "Conclusion", "Future Works"
    ]
    for sec in order:
        content = sections.get(sec, "").strip()
        if not content:
            continue
        doc.add_heading(sec, level=1)
        for para in re.split(r"\n\s*\n", content):
            if para.strip():
                doc.add_paragraph(para.strip())

    if references_md.strip():
        doc.add_heading("References", level=1)
        for line in references_md.splitlines():
            if line.strip():
                doc.add_paragraph(line)

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)


def write_pdf(path: str, title: str, sections: Dict[str, str], references_md: str):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    c = canvas.Canvas(path, pagesize=LETTER)
    width, height = LETTER
    margin = 0.75 * inch
    x, y = margin, height - margin

    def draw_line(text: str, bold: bool = False, font_size: int = 10, line_spacing: int = 14):
        nonlocal x, y
        if y < margin:
            c.showPage()
            y = height - margin
        font_name = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font_name, font_size)
        words = text.split()
        line = ""
        for w in words:
            candidate = (line + " " + w).strip()
            if c.stringWidth(candidate, font_name, font_size) > (width - 2 * margin):
                c.drawString(x, y, line)
                y -= line_spacing
                if y < margin:
                    c.showPage()
                    y = height - margin
                line = w
            else:
                line = candidate
        if line:
            c.drawString(x, y, line)
            y -= line_spacing

    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(x, y, title)
    y -= 24

    order = [
        "Abstract", "Introduction", "Literature Review",
        "Key Findings", "Results & Discussion", "Conclusion", "Future Works"
    ]
    for sec in order:
        content = sections.get(sec, "").strip()
        if not content:
            continue
        draw_line(sec, bold=True, font_size=12, line_spacing=16)
        for para in re.split(r"\n\s*\n", content):
            if para.strip():
                draw_line(para, bold=False, font_size=10, line_spacing=14)
        y -= 6

    if references_md.strip():
        draw_line("References", bold=True, font_size=12, line_spacing=16)
        for line in references_md.splitlines():
            if line.strip():
                draw_line(line, bold=False, font_size=10, line_spacing=14)

    c.save()
